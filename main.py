from docx import Document
from docx.shared import RGBColor, Pt

def encode_phrase_and_modify_doc_color(phrase, doc_path, output_doc_path):
    # Кодирование фразы в cp866 и преобразование в двоичный формат
    encoded_phrase = phrase.encode("cp866")
    binary_phrase = ''.join(format(byte, '08b') for byte in encoded_phrase)
    print(binary_phrase)

    # Открытие документа
    doc = Document(doc_path)

    encoding_started = False
    binary_index = 0

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size == Pt(15) or encoding_started:
                encoding_started = True
                new_text = ''
                for char in run.text:
                    new_run = paragraph.add_run(char)
                    new_run.font.name = 'Segoe Print'  # Установка шрифта Georgia
                    if binary_index < len(binary_phrase):
                        new_run.font.size = Pt(16 if binary_phrase[binary_index] == '1' else 15)
                        binary_index += 1
                    else:
                        new_run.font.size = Pt(15)
                    new_run.font.color.rgb = run.font.color.rgb if run.font.color else RGBColor(0,0,0)
                    new_text += char
                run.text = ''  # Очищаем исходный текст

    # Сохранение изменений в новом документе
    doc.save(output_doc_path)

# Пример использования функции
phrase = "Без учебы и труда не придет на стол еда."
doc_path = "8.docx"  # Укажите путь к исходному документу
output_doc_path = "../pythonProject5/Итог.docx"  # Укажите путь к итоговому документу
encode_phrase_and_modify_doc_color(phrase, doc_path, output_doc_path)
